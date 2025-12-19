import os
# Suppress TensorFlow and oneDNN warnings
os.environ["TF_ENABLE_ONEDNN_OPTS"] = "0"
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "2"
# Allow OAuth over HTTP for local development
os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"

from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, RedirectResponse, HTMLResponse
from groq import Groq
from langchain_text_splitters import RecursiveCharacterTextSplitter
# Try to import from new langchain-huggingface package, fall back to deprecated location
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.vectorstores.utils import DistanceStrategy
from langchain_community.document_loaders import UnstructuredExcelLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
import pdfplumber
import requests

import re
import tempfile
import shutil
import glob
import time
import threading
from pathlib import Path
from dotenv import load_dotenv
from pydantic import BaseModel
from typing import List, Optional
import traceback
import io
import base64
import json
from urllib.parse import urlparse, parse_qs
import pandas as pd

# Google OAuth and Drive imports
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# Import from local modules
from text_extraction import (
    extract_text_by_mimetype, 
    extract_text_from_file,
    extract_pdf_with_tables,
    extract_docx_with_tables,
    is_tesseract_available as extraction_tesseract_available,
    TESSERACT_AVAILABLE as EXTRACTION_TESSERACT_AVAILABLE
)
from whatsapp import WhatsAppScraper, is_tesseract_available as whatsapp_tesseract_available

# Import MCP client for Google Drive and Gmail integration (unified client)
from mcp_client import get_mcp_client, MCPDriveClient

# Note: OCR imports are now in text_extraction.py
# Image import is still needed for some direct operations
from PIL import Image

# LangChain Agent imports for Excel processing and Agentic Router
try:
    from langchain_groq import ChatGroq
    from langchain_experimental.agents.agent_toolkits import create_pandas_dataframe_agent
    from langchain.agents import Tool, initialize_agent, AgentType
    from langchain.chains import RetrievalQA
    LANGCHAIN_AGENT_AVAILABLE = True
except ImportError:
    LANGCHAIN_AGENT_AVAILABLE = False
    print("[WARNING] LangChain agent dependencies not available. Excel agent will be disabled.")


load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global variables
UPLOAD_FOLDER = "uploads"
WHATSAPP_DOWNLOAD_DIR = os.path.join(os.getcwd(), "whatsapp_downloads")
WHATSAPP_IMAGES_DIR = os.path.join(os.getcwd(), "whatsapp_images")
GOOGLE_DRIVE_DOWNLOAD_DIR = os.path.join(os.getcwd(), "google_drive_downloads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(WHATSAPP_DOWNLOAD_DIR, exist_ok=True)
os.makedirs(WHATSAPP_IMAGES_DIR, exist_ok=True)
os.makedirs(GOOGLE_DRIVE_DOWNLOAD_DIR, exist_ok=True)

# Google OAuth Configuration
CLIENT_ID = os.getenv('CLIENT_ID')
CLIENT_SECRET = os.getenv('CLIENT_SECRET')
REDIRECT_URI = os.getenv('REDIRECT_URI', 'http://localhost:5000/oauth2callback')
SCOPES = [
    'https://www.googleapis.com/auth/drive.readonly',
    'https://www.googleapis.com/auth/gmail.readonly',  # Gmail read access for email extraction
    'https://www.googleapis.com/auth/gmail.send',  # Gmail send access for replying to emails
]

# Store credentials temporarily (use database in production)
credentials_store = {}

# Global state
whatsapp_driver = None
whatsapp_connected = False
rag_system = None
excel_agent_system = None  # Excel agent for quantitative queries
agentic_router = None  # Intelligent agentic router for tool selection
selected_groups: List[str] = []

# Thread locks for concurrent access safety
_rag_lock = threading.Lock()  # Lock for RAG system / vector store operations
_excel_lock = threading.Lock()  # Lock for Excel agent operations

# Global embeddings instance - pre-loaded at startup for faster first query
_global_embeddings = None
_embeddings_lock = threading.Lock()
_embeddings_loading = False


def _get_or_create_embeddings():
    """Get the global embeddings instance, creating it if necessary (thread-safe)."""
    global _global_embeddings, _embeddings_loading
    
    # Fast path: already loaded
    if _global_embeddings is not None:
        return _global_embeddings
    
    with _embeddings_lock:
        # Double-check after acquiring lock
        if _global_embeddings is not None:
            return _global_embeddings
        
        if _embeddings_loading:
            # Another thread is loading, wait for it
            print("[INFO] Waiting for embeddings to be loaded by another thread...")
            return None
        
        _embeddings_loading = True
        print("[INFO] Loading embeddings model...")
        try:
            _global_embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={"device": "cpu"},
                encode_kwargs={"normalize_embeddings": True},
            )
            print("[INFO] Embeddings model loaded successfully!")
        except Exception as e:
            print(f"[ERROR] Failed to load embeddings: {e}")
            _embeddings_loading = False
            raise
        finally:
            _embeddings_loading = False
        
        return _global_embeddings


# Startup event - Reset vector store when backend restarts
@app.on_event("startup")
async def startup_event():
    """Reset all in-memory systems on startup and pre-load embeddings."""
    global rag_system, excel_agent_system, agentic_router
    print("[INFO] ========================================")
    print("[INFO] TOAI Backend Starting...")
    print("[INFO] Resetting vector store and agents...")
    rag_system = None
    excel_agent_system = None
    agentic_router = None
    print("[INFO] Vector store and agents cleared.")
    
    # Pre-load embeddings synchronously at startup
    try:
        _get_or_create_embeddings()
    except Exception as e:
        print(f"[WARNING] Could not pre-load embeddings: {e}")
    
    print("[INFO] Ready to accept connections!")
    print("[INFO] ========================================")



def format_tables_in_response(text: str) -> str:
    """
    Detect and convert plain-text tables in the response to properly formatted markdown tables.
    This handles tables output by pandas print() statements that aren't using to_markdown().
    
    Detects patterns like:
        column1     column2     column3
    row1    1.234       5.678       9.012
    row2    2.345       6.789       0.123
    
    And converts to:
    | row | column1 | column2 | column3 |
    |:----|--------:|--------:|--------:|
    | row1 | 1.234 | 5.678 | 9.012 |
    | row2 | 2.345 | 6.789 | 0.123 |
    """
    if not text:
        return text
    
    # If already contains markdown table syntax, return as-is
    if '|' in text and '---' in text:
        return text
    
    lines = text.split('\n')
    result_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Detect potential table start: line with multiple whitespace-separated items
        # that looks like a header (multiple words/columns separated by 2+ spaces)
        if re.match(r'^[\w_]+(\s{2,}[\w_%.]+){2,}', line.strip()) or \
           re.match(r'^(\s+[\w_%.]+){2,}', line):
            # Try to parse as a table
            table_lines = []
            table_start = i
            
            # Collect consecutive lines that look like table data
            while i < len(lines):
                current_line = lines[i]
                
                # Empty line or line that doesn't fit table pattern ends the table
                if not current_line.strip():
                    break
                
                # Check if line has multiple columns (separated by 2+ spaces)
                parts = re.split(r'\s{2,}', current_line.strip())
                if len(parts) >= 2:
                    table_lines.append(current_line)
                    i += 1
                else:
                    break
            
            # If we collected enough lines for a table (header + at least 1 data row)
            if len(table_lines) >= 2:
                try:
                    markdown_table = _convert_text_table_to_markdown(table_lines)
                    if markdown_table:
                        result_lines.append(markdown_table)
                        continue
                except Exception as e:
                    print(f"[DEBUG] Table conversion failed: {e}")
                    # Fall through to add lines as-is
            
            # If table parsing failed, add lines as-is
            for j in range(table_start, i):
                result_lines.append(lines[j])
        else:
            result_lines.append(line)
            i += 1
    
    return '\n'.join(result_lines)


def _convert_text_table_to_markdown(table_lines: list) -> str:
    """
    Convert a list of text table lines to a markdown table.
    """
    if not table_lines:
        return ""
    
    # Parse each line into columns
    parsed_rows = []
    max_cols = 0
    
    for line in table_lines:
        # Split by 2+ whitespace characters
        parts = re.split(r'\s{2,}', line.strip())
        parts = [p.strip() for p in parts if p.strip()]
        if parts:
            parsed_rows.append(parts)
            max_cols = max(max_cols, len(parts))
    
    if not parsed_rows or max_cols < 2:
        return ""
    
    # Normalize all rows to have the same number of columns
    for row in parsed_rows:
        while len(row) < max_cols:
            row.append("")
    
    # Check if first column seems to be an index (row labels)
    first_row = parsed_rows[0]
    has_index = False
    
    # If the first row's first cell is empty or a known index label, it's an index column
    if first_row[0] == '' or first_row[0].lower() in ['', 'index', 'city', 'name', 'id', 'row']:
        has_index = True
    
    # Build markdown table
    md_lines = []
    
    # Header row
    header = parsed_rows[0]
    if has_index and header[0] == '':
        header[0] = 'Index'
    md_lines.append('| ' + ' | '.join(header) + ' |')
    
    # Separator row (right-align numeric columns)
    separators = []
    for col_idx, col in enumerate(header):
        # Check if this column has numeric values
        is_numeric = True
        for row in parsed_rows[1:]:
            if col_idx < len(row):
                try:
                    float(row[col_idx].replace('%', '').replace(',', ''))
                except ValueError:
                    is_numeric = False
                    break
        
        if is_numeric:
            separators.append('---:')  # Right-align numbers
        else:
            separators.append(':---')  # Left-align text
    
    md_lines.append('| ' + ' | '.join(separators) + ' |')
    
    # Data rows
    for row in parsed_rows[1:]:
        md_lines.append('| ' + ' | '.join(row) + ' |')
    
    return '\n'.join(md_lines)


class ExcelAgentSystem:
    """
    System for handling Excel-specific queries using a pandas dataframe agent.
    Routes quantitative questions (aggregations, filtering, counts) to the Excel agent,
    while qualitative questions go to the RAG system.
    """
    
    # Keywords that indicate a query should be routed to the Excel agent
    EXCEL_KEYWORDS = [
        'how many', 'count', 'total', 'sum', 'average', 'mean', 'median',
        'maximum', 'minimum', 'max', 'min', 'filter', 'where', 'group by',
        'sort', 'rank', 'top', 'bottom', 'percentage', 'percent', '%',
        'calculate', 'compute', 'aggregate', 'distinct', 'unique',
        'greater than', 'less than', 'between', 'equal to', 'not equal',
        'rows', 'columns', 'cells', 'values', 'data points',
        'excel', 'spreadsheet', 'sheet', 'worksheet'
    ]
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.dataframes: dict = {}  # {filename: DataFrame}
        self.agents: dict = {}  # {filename: pandas_agent}
        self.llm = None
        if LANGCHAIN_AGENT_AVAILABLE:
            self.llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0)
    
    def add_excel_file(self, file_path: str) -> bool:
        """Load an Excel file and create a pandas agent for it. Thread-safe for concurrent uploads."""
        if not LANGCHAIN_AGENT_AVAILABLE:
            print("[WARNING] LangChain agent not available. Cannot process Excel file with agent.")
            return False
        
        # Use lock for thread-safe Excel agent operations
        with _excel_lock:
            try:
                filename = os.path.basename(file_path)
                
                # Read Excel file - handle multiple sheets by concatenating
                excel_data = pd.read_excel(file_path, sheet_name=None)
                
                if len(excel_data) == 1:
                    # Single sheet
                    df = list(excel_data.values())[0]
                else:
                    # Multiple sheets - add sheet name as column and concatenate
                    dfs = []
                    for sheet_name, sheet_df in excel_data.items():
                        sheet_df = sheet_df.copy()
                        sheet_df['_sheet_name'] = sheet_name
                        dfs.append(sheet_df)
                    df = pd.concat(dfs, ignore_index=True)
                
                self.dataframes[filename] = df
                
                # Custom prefix to clarify the agent has access to the ENTIRE dataframe
                # This fixes the issue where the LLM incorrectly assumes only the sample head is available
                custom_prefix = f"""You are working with a pandas dataframe in Python. The dataframe name is `df`.

IMPORTANT RULES:
1. You have access to the ENTIRE dataframe with {len(df)} rows and {len(df.columns)} columns.
2. The sample data shown below is ONLY for reference to understand the structure. 
3. You MUST execute Python code to answer questions - NEVER guess based on the sample.
4. Always use the full dataframe `df` when computing counts, aggregations, or any analysis.

CRITICAL - ACTION FORMAT:
- You can ONLY use the tool named: python_repl_ast
- Your Action must ALWAYS be exactly: python_repl_ast
- NEVER use any other action name like "Filter the dataframe" or similar
- Format your action as:
  Action: python_repl_ast
  Action Input: <your python code here>

OUTPUT RULES:
- When asked for a list of items (like product names), ALWAYS print the actual values, not just counts
- Use print() to show the results clearly
- For lists, print the actual titles/names, not just "there are X items"
- Set pandas display options to show full output: pd.set_option('display.max_rows', None)
- When printing dataframes, select only the relevant columns to show (e.g., df[['TITLE', 'Rating']])

TABLE FORMATTING (CRITICAL):
- ALWAYS format table/dataframe output as MARKDOWN TABLES using .to_markdown()
- Example: print(result_df.to_markdown(index=True))
- For grouped results: print(df.groupby('city')['column'].mean().to_frame().to_markdown())
- This ensures tables are properly formatted and readable
- NEVER use .to_string() or plain print(df) for tables - ALWAYS use .to_markdown()

EXAMPLES:
- To count unique values: print(df['column'].nunique())
- To get list of items: print(df['TITLE'].tolist())
- To filter and show: print(df[df['Rating'] == 5]['TITLE'].tolist())
- To show table with multiple columns: print(df[df['Rating'] == 5][['TITLE', 'Rating']].to_markdown(index=False))
- To show grouped percentages: print(city_percentages.to_markdown())

Column names: {', '.join(df.columns.tolist())}
"""
                
                # Create a pandas agent for this dataframe with custom prefix
                pandas_agent = create_pandas_dataframe_agent(
                    self.llm,
                    df,
                    verbose=True,
                    allow_dangerous_code=True,
                    handle_parsing_errors=True,
                    prefix=custom_prefix,
                    max_iterations=10,
                    early_stopping_method="generate"
                )
                
                self.agents[filename] = pandas_agent
                
                print(f"[INFO] Excel agent created for: {filename} ({len(df)} rows, {len(df.columns)} columns)")
                return True
                
            except Exception as e:
                print(f"[ERROR] Failed to create Excel agent for {file_path}: {e}")
                traceback.print_exc()
                return False
    
    def is_excel_query(self, query: str) -> bool:
        """Determine if a query should be routed to the Excel agent."""
        query_lower = query.lower()
        
        # Check for Excel-related keywords
        for keyword in self.EXCEL_KEYWORDS:
            if keyword in query_lower:
                return True
        
        # Check if query mentions specific Excel file names
        for filename in self.dataframes.keys():
            if filename.lower().replace('.xlsx', '').replace('.xls', '') in query_lower:
                return True
        
        return False
    
    def run_query(self, query: str) -> str:
        """Run a query against the Excel data using the pandas agent."""
        if not LANGCHAIN_AGENT_AVAILABLE:
            return "Excel agent is not available. Please install langchain-groq and langchain-experimental."
        
        if not self.agents:
            return "No Excel files have been loaded. Please upload an Excel file first."
        
        try:
            # If we have multiple files, we need to determine which one the user is asking about
            # For simplicity, we'll use the first (or most recently added) agent
            # You could enhance this with more sophisticated file selection logic
            
            if len(self.agents) == 1:
                agent = list(self.agents.values())[0]
                filename = list(self.agents.keys())[0]
            else:
                # Try to detect which file the query is about
                agent = None
                filename = None
                for fname in self.agents.keys():
                    if fname.lower().replace('.xlsx', '').replace('.xls', '') in query.lower():
                        agent = self.agents[fname]
                        filename = fname
                        break
                
                # If no specific file mentioned, use the most recent one (last added)
                if agent is None:
                    filename = list(self.agents.keys())[-1]
                    agent = self.agents[filename]
            
            print(f"[INFO] Running Excel agent query on: {filename}")
            
            # Run the query through the pandas agent
            result = agent.run(query)
            
            return f"[Analysis from {filename}]\n{result}"
            
        except Exception as e:
            print(f"[ERROR] Excel agent query failed: {e}")
            traceback.print_exc()
            return f"Error analyzing Excel data: {str(e)}"
    
    def get_excel_context_for_rag(self) -> str:
        """Generate a summary context of Excel data for the RAG system."""
        if not self.dataframes:
            return ""
        
        context_parts = []
        for filename, df in self.dataframes.items():
            summary = f"\n[Excel File: {filename}]\n"
            summary += f"Columns: {', '.join(df.columns.tolist())}\n"
            summary += f"Total Rows: {len(df)}\n"
            
            # Add sample data (first 3 rows)
            summary += "Sample Data:\n"
            summary += df.head(3).to_string(index=False) + "\n"
            
            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
            if numeric_cols:
                summary += f"Numeric columns: {', '.join(numeric_cols)}\n"
            
            context_parts.append(summary)
        
        return "\n".join(context_parts)
    
    def create_rag_documents(self) -> list:
        """
        Create LangChain Document objects from Excel data for the RAG system.
        This feeds Excel structure and sample data into the vector store.
        """
        documents = []
        
        for filename, df in self.dataframes.items():
            # Create a comprehensive summary document for each Excel file
            content_parts = []
            
            # Basic info
            content_parts.append(f"Excel File: {filename}")
            content_parts.append(f"Total Rows: {len(df)}")
            content_parts.append(f"Total Columns: {len(df.columns)}")
            content_parts.append(f"Column Names: {', '.join(df.columns.tolist())}")
            
            # Column types
            content_parts.append("\nColumn Types:")
            for col in df.columns:
                dtype = str(df[col].dtype)
                content_parts.append(f"  - {col}: {dtype}")
            
            # Sample data (first 5 rows)
            content_parts.append("\nSample Data (first 5 rows):")
            content_parts.append(df.head(5).to_string(index=False))
            
            # Basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
            if numeric_cols:
                content_parts.append("\nNumeric Column Statistics:")
                for col in numeric_cols:
                    content_parts.append(f"  {col}:")
                    content_parts.append(f"    - Min: {df[col].min()}")
                    content_parts.append(f"    - Max: {df[col].max()}")
                    content_parts.append(f"    - Mean: {df[col].mean():.2f}")
            
            # Unique value counts for categorical columns (first 10 unique values)
            categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
            if categorical_cols:
                content_parts.append("\nCategorical Column Unique Values:")
                for col in categorical_cols[:5]:  # Limit to first 5 categorical columns
                    unique_count = df[col].nunique()
                    content_parts.append(f"  {col}: {unique_count} unique values")
                    if unique_count <= 10:
                        unique_vals = df[col].dropna().unique().tolist()[:10]
                        content_parts.append(f"    Values: {unique_vals}")
            
            content = "\n".join(content_parts)
            
            doc = Document(
                page_content=content,
                metadata={
                    "source": filename,
                    "file_type": "excel_summary",
                    "rows": len(df),
                    "columns": len(df.columns)
                }
            )
            documents.append(doc)
            print(f"[INFO] Created RAG document from Excel file: {filename}")
        
        return documents
    
    def clear(self):
        """Clear all loaded Excel data and agents."""
        self.dataframes.clear()
        self.agents.clear()


class AgenticRouter:
    """
    An intelligent agentic router that uses an LLM to decide which tool(s) to use.
    The agent can route queries to:
    - PDF/Document Knowledge Base (RAG) for qualitative questions
    - Excel Data Analyst for quantitative questions  
    - Both tools when needed for comprehensive answers
    """
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.llm = None
        self.agent = None
        self.tools = []
        self.rag_tool = None
        self.excel_tool = None
        self._initialized = False
        
        if LANGCHAIN_AGENT_AVAILABLE:
            self.llm = ChatGroq(
                model_name="llama-3.3-70b-versatile", 
                temperature=0,
                api_key=api_key
            )
    
    def setup_rag_tool(self, rag_system):
        """
        Setup the RAG tool for querying PDF/document knowledge base.
        """
        if not LANGCHAIN_AGENT_AVAILABLE:
            print("[WARNING] LangChain not available. Cannot setup RAG tool.")
            return None
            
        if not rag_system or not rag_system.vector_store:
            print("[WARNING] RAG system or vector store not available.")
            return None
        
        try:
            # Create a RetrievalQA chain for the RAG system
            qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=rag_system.vector_store.as_retriever(
                    search_kwargs={"k": 5}
                )
            )
            
            self.rag_tool = Tool(
                name="PDF_Document_Knowledge_Base",
                func=qa_chain.run,
                description=(
                    "Useful for answering qualitative questions, summaries, "
                    "checking policies, finding information from PDF documents, "
                    "Word documents, and other text-based files. Use this tool for "
                    "questions about document content, procedures, guidelines, "
                    "descriptions, or any non-numerical analysis."
                ),
                return_direct=True  # Return tool output directly without further agent reasoning
            )
            
            print("[INFO] RAG tool setup complete")
            return self.rag_tool
            
        except Exception as e:
            print(f"[ERROR] Failed to setup RAG tool: {e}")
            traceback.print_exc()
            return None
    
    def setup_excel_tool(self, excel_agent_system):
        """
        Setup the Excel tool for querying spreadsheet data.
        """
        if not LANGCHAIN_AGENT_AVAILABLE:
            print("[WARNING] LangChain not available. Cannot setup Excel tool.")
            return None
            
        if not excel_agent_system or not excel_agent_system.agents:
            print("[WARNING] Excel agent system not available or no Excel files loaded.")
            return None
        
        try:
            # Create a wrapper function that routes to the excel agent
            def excel_query_wrapper(query: str) -> str:
                return excel_agent_system.run_query(query)
            
            self.excel_tool = Tool(
                name="Excel_Data_Analyst",
                func=excel_query_wrapper,
                description=(
                    "Useful for quantitative questions, mathematical calculations, "
                    "aggregations, filtering data, counting records, or 'how many' "
                    "questions about Excel/spreadsheet data. Use this tool for "
                    "questions involving numbers, totals, averages, counts, "
                    "min/max values, percentages, or data analysis."
                ),
                return_direct=True  # Return tool output directly without further agent reasoning
            )
            
            print("[INFO] Excel tool setup complete")
            return self.excel_tool
            
        except Exception as e:
            print(f"[ERROR] Failed to setup Excel tool: {e}")
            traceback.print_exc()
            return None
    
    def initialize_agent(self, rag_system=None, excel_agent_system=None):
        """
        Initialize the agentic router with available tools.
        Must be called after tools are setup.
        """
        if not LANGCHAIN_AGENT_AVAILABLE:
            print("[WARNING] Cannot initialize agent - LangChain not available")
            return False
        
        self.tools = []
        
        # Setup RAG tool if RAG system is available
        if rag_system and rag_system.vector_store:
            rag_tool = self.setup_rag_tool(rag_system)
            if rag_tool:
                self.tools.append(rag_tool)
        
        # Setup Excel tool if Excel agent is available
        if excel_agent_system and excel_agent_system.agents:
            excel_tool = self.setup_excel_tool(excel_agent_system)
            if excel_tool:
                self.tools.append(excel_tool)
        
        if not self.tools:
            print("[WARNING] No tools available for agent")
            return False
        
        try:
            # Initialize the main "brain" agent
            self.agent = initialize_agent(
                self.tools,
                self.llm,
                agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
                verbose=True,
                handle_parsing_errors=True,
                max_iterations=5,
                early_stopping_method="generate"
            )
            
            self._initialized = True
            tool_names = [t.name for t in self.tools]
            print(f"[INFO] Agentic router initialized with tools: {tool_names}")
            return True
            
        except Exception as e:
            print(f"[ERROR] Failed to initialize agent: {e}")
            traceback.print_exc()
            return False
    
    def run(self, query: str) -> dict:
        """
        Run a query through the agentic router.
        The agent will intelligently decide which tool(s) to use.
        
        Returns:
            dict with 'response', 'tools_used', and 'success' keys
        """
        if not LANGCHAIN_AGENT_AVAILABLE:
            return {
                "response": "Agentic router not available. LangChain dependencies missing.",
                "tools_used": [],
                "success": False
            }
        
        if not self._initialized or not self.agent:
            return {
                "response": "Agentic router not initialized. Please upload documents first.",
                "tools_used": [],
                "success": False
            }
        
        try:
            print(f"\n{'='*60}")
            print(f"[AGENTIC ROUTER] Processing query: {query[:80]}...")
            print(f"[AGENTIC ROUTER] Available tools: {[t.name for t in self.tools]}")
            print(f"{'='*60}")
            
            # Run the agent
            response = self.agent.run(query)
            
            # Determine which tools were used (from agent's verbose output)
            tools_used = []
            if self.rag_tool and "PDF_Document_Knowledge_Base" in str(self.agent):
                tools_used.append("PDF_Document_Knowledge_Base")
            if self.excel_tool and "Excel_Data_Analyst" in str(self.agent):
                tools_used.append("Excel_Data_Analyst")
            
            print(f"[AGENTIC ROUTER] Response generated successfully")
            
            return {
                "response": response,
                "tools_used": tools_used,
                "success": True
            }
            
        except Exception as e:
            print(f"[ERROR] Agentic router failed: {e}")
            traceback.print_exc()
            return {
                "response": f"Error in agentic router: {str(e)}",
                "tools_used": [],
                "success": False
            }
    
    def is_ready(self) -> bool:
        """Check if the agentic router is ready to process queries."""
        return self._initialized and self.agent is not None and len(self.tools) > 0
    
    def get_available_tools(self) -> list:
        """Get list of available tool names."""
        return [t.name for t in self.tools] if self.tools else []


class RAGSystem:
    def __init__(self, api_key, embeddings=None):
        self.groq_client = Groq(api_key=api_key)
        if embeddings is not None:
            self.embeddings = embeddings
        else:
            # Use thread-safe embeddings getter
            self.embeddings = _get_or_create_embeddings()
            if self.embeddings is None:
                # Embeddings still loading, create a local one
                print("[INFO] Creating local embeddings instance...")
                self.embeddings = HuggingFaceEmbeddings(
                    model_name="sentence-transformers/all-MiniLM-L6-v2",
                    model_kwargs={"device": "cpu"},
                    encode_kwargs={"normalize_embeddings": True},
                )
            else:
                print("[INFO] Using pre-loaded embeddings")
        self.vector_store = None
        self.chat_history = []

    def load_documents(self, file_paths):
        """
        Load documents from file paths using centralized text extraction.
        Uses extract_text_from_file from text_extraction module.
        """
        documents = []

        for file_path in file_paths:
            try:
                file_name = os.path.basename(file_path)
                suffix = Path(file_name).suffix.lower()
                
                # Skip Excel files - they are processed exclusively by the Excel Agent
                if suffix in (".xlsx", ".xls"):
                    print(f"[INFO] Skipping {file_name} in RAG - Excel files are processed by Excel Agent only")
                    continue
                
                # Use centralized text extraction
                content, file_type, error = extract_text_from_file(file_path)
                
                if error:
                    print(f"[WARNING] {error}")
                    continue
                
                if content and content.strip():
                    # For images, add prefix
                    if file_type == "ocr_image":
                        content = f"[Image: {file_name}]\n{content}"
                        print(f"✓ OCR text extracted from {file_name} ({len(content)} characters)")
                    
                    doc = Document(
                        page_content=content,
                        metadata={"source": file_name, "file_type": file_type},
                    )
                    documents.append(doc)
                else:
                    print(f"[INFO] No text content extracted from {file_name}")
                    
            except Exception as e:
                print(f"Error loading {file_path}: {e}")
                continue

        return documents

    def load_whatsapp_messages(self, messages):
        """Load WhatsApp messages as documents"""
        documents = []

        for msg in messages:
            content = f"[{msg['timestamp']}] {msg['sender']}: {msg['text']}"
            doc = Document(
                page_content=content,
                metadata={"source": "WhatsApp", "file_type": "chat"},
            )
            documents.append(doc)

        return documents

    def load_gmail_emails(self, emails):
        """Load Gmail emails as documents for RAG system
        
        Args:
            emails: List of email dictionaries with 'subject', 'from', 'date', 'content' keys
        
        Returns:
            List of Document objects
        """
        documents = []

        for email in emails:
            subject = email.get('subject', '(No Subject)')
            sender = email.get('from', 'Unknown Sender')
            date = email.get('date', 'Unknown Date')
            content = email.get('content', '') or email.get('fullText', '') or email.get('body', '')
            
            # Create formatted content for the document
            formatted_content = f"""[Email] Subject: {subject}
From: {sender}
Date: {date}

{content}"""
            
            doc = Document(
                page_content=formatted_content,
                metadata={
                    "source": "Gmail",
                    "file_type": "email",
                    "email_id": email.get('id', ''),
                    "subject": subject,
                    "sender": sender
                },
            )
            documents.append(doc)
            print(f"[RAG] Loaded email: {subject[:50]}..." if len(subject) > 50 else f"[RAG] Loaded email: {subject}")

        return documents

    def load_ocr_texts(self, ocr_results):
        """Load OCR extracted text from images as documents"""
        documents = []

        for ocr_item in ocr_results:
            image_path = ocr_item.get('image', 'Unknown')
            text = ocr_item.get('text', '')
            item_type = ocr_item.get('type', 'image')
            
            if text.strip():
                image_name = os.path.basename(image_path)
                content = f"[Image: {image_name}]\n{text}"
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": "WhatsApp",
                        "file_type": "ocr_image",
                        "image_path": image_path,
                        "item_type": item_type
                    },
                )
                documents.append(doc)

        return documents

    def _merge_and_split_sequential_tables(self, documents):
        processed_docs = []

        for doc in documents:
            content = doc.page_content
            metadata = doc.metadata.copy()
            table_pattern = r"----- TABLE -----\n(.*?)\n----- END TABLE -----"
            parts = []
            last_end = 0
            table_matches = list(re.finditer(table_pattern, content, re.DOTALL))

            if not table_matches:
                processed_docs.append(doc)
                continue

            for i, match in enumerate(table_matches):
                text_before = content[last_end : match.start()].strip()
                text_without_whitespace = re.sub(r"\s+", "", text_before)
                if text_without_whitespace:
                    parts.append(("text", text_before))

                table_content = match.group(1)
                table_rows = [row for row in table_content.split("\n") if row.strip()]
                parts.append(("table", table_rows))
                last_end = match.end()

            text_after = content[last_end:].strip()
            text_after_without_whitespace = re.sub(r"\s+", "", text_after)
            if text_after_without_whitespace:
                parts.append(("text", text_after))

            final_parts = []
            sequential_table_rows = []
            MAX_ROWS_PER_CHUNK = 50

            for part_type, part_content in parts:
                if part_type == "text":
                    text_without_whitespace = re.sub(r"\s+", "", part_content)
                    if text_without_whitespace:
                        if sequential_table_rows:
                            for i in range(
                                0, len(sequential_table_rows), MAX_ROWS_PER_CHUNK
                            ):
                                chunk_rows = sequential_table_rows[
                                    i : i + MAX_ROWS_PER_CHUNK
                                ]
                                chunk_table_text = self._format_table_group(
                                    [chunk_rows]
                                )
                                final_parts.append(("table_chunk", chunk_table_text))
                            sequential_table_rows = []
                        final_parts.append(("text", part_content))
                elif part_type == "table":
                    sequential_table_rows.extend(part_content)

            if sequential_table_rows:
                for i in range(0, len(sequential_table_rows), MAX_ROWS_PER_CHUNK):
                    chunk_rows = sequential_table_rows[i : i + MAX_ROWS_PER_CHUNK]
                    chunk_table_text = self._format_table_group([chunk_rows])
                    final_parts.append(("table_chunk", chunk_table_text))

            reconstructed_parts = []
            for part_type, part_content in final_parts:
                if part_type in ["text", "table_chunk"]:
                    reconstructed_parts.append(part_content)

            new_content = "\n\n".join(reconstructed_parts)
            processed_doc = Document(page_content=new_content, metadata=metadata)
            processed_docs.append(processed_doc)

        return processed_docs

    def _format_table_group(self, table_group):
        all_rows = []
        for table_rows in table_group:
            all_rows.extend(table_rows)
        table_text = "\n".join(all_rows)
        return f"----- TABLE -----\n{table_text}\n----- END TABLE -----"

    def create_vector_store(self, documents):
        """Create or update vector store with documents. Thread-safe for concurrent uploads."""
        processed_documents = self._merge_and_split_sequential_tables(documents)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200,
        )

        all_chunks = []
        for doc in processed_documents:
            splits = text_splitter.split_documents([doc])
            all_chunks.extend(splits)
        print(f"[DEBUG] Preparing to index {len(all_chunks)} chunks from {len(documents)} documents.")

        if not all_chunks:
            print("[ERROR] No valid chunks found in documents, nothing to index!")
            raise ValueError("No valid chunks found in documents.")

        # Use lock for thread-safe vector store operations
        with _rag_lock:
            if self.vector_store is None:
                print("[DEBUG] Initializing new FAISS vector store.")
                self.vector_store = FAISS.from_documents(
                    all_chunks,
                    self.embeddings,
                    distance_strategy=DistanceStrategy.COSINE,
                    normalize_L2=True,
                )
            else:
                print(f"[DEBUG] Adding {len(all_chunks)} chunks to existing vector store.")
                # Add to existing vector store
                self.vector_store.add_documents(all_chunks)

            print(f"[DEBUG] Vector store now contains documents from these sources: ")
            # Attempt to print short list of sources for inspection
            try:
                if hasattr(self.vector_store, 'docs'):
                    sources = [doc.metadata.get("source", "Unknown") for doc in self.vector_store.docs]
                    print(f"[DEBUG] Current sources: {sources}")
            except Exception as e:
                print(f"[DEBUG] Unable to enumerate sources in vector store: {e}")

        return len(all_chunks)

    def retrieve_context(self, query, k=5):
        """Retrieve context from vector store. Thread-safe for concurrent access.
        Returns tuple of (context_string, list_of_sources)
        """
        with _rag_lock:
            if self.vector_store is None:
                return "", []
            docs = self.vector_store.similarity_search(query, k=k)
        
        # Build context with source attribution
        context_parts = []
        sources = []
        for doc in docs:
            source = doc.metadata.get('source', 'Unknown')
            file_type = doc.metadata.get('file_type', 'document')
            
            # Track unique sources
            if source not in sources:
                sources.append(source)
            
            context_parts.append(f"Source: {source}\n{doc.page_content}")
        
        context = "\n\n---\n\n".join(context_parts)
        return context, sources

    def generate_response(self, query, context, sources=None):
        """Generate response with citations from the sources."""
        if sources is None:
            sources = []
        
        # Build sources citation hint
        sources_hint = ""
        if sources:
            sources_list = ", ".join(sources)
            sources_hint = f"\n\nAvailable sources for citation: {sources_list}"
        
        messages = [
            {
                "role": "system",
                "content": """You are a helpful assistant that answers questions based on documents and WhatsApp conversations. 
                
IMPORTANT: When answering questions from the provided context, you MUST:
1. Cite the source document(s) you used to form your answer
2. At the end of your response, include a "Sources:" section listing all referenced documents
3. Use the exact source names provided in the context (e.g., "Source: filename.pdf")
4. If information comes from WhatsApp, cite it as "WhatsApp messages"
5. If information comes from Gmail, cite it as "Gmail - [subject]"

Example format:
[Your answer here]

**Sources:**
- document1.pdf
- document2.pdf
- WhatsApp messages

Maintain context from previous questions.""",
            }
        ]

        for msg in self.chat_history[-5:]:
            messages.append(msg)

        prompt = f"""Based on the following context, please answer the question. Remember to cite your sources.

Context:
{context}
{sources_hint}

Question: {query}

Provide a detailed answer based on the context and include citations."""

        messages.append({"role": "user", "content": prompt})

        try:
            chat_completion = self.groq_client.chat.completions.create(
                messages=messages,
                model="llama-3.3-70b-versatile",
                temperature=0.3,
                max_tokens=2048,
            )

            response = chat_completion.choices[0].message.content
            self.chat_history.append({"role": "user", "content": query})
            self.chat_history.append({"role": "assistant", "content": response})
            return response
        except Exception as e:
            return f"Error: {str(e)}"


# Note: Text extraction functions (PDF, DOCX, images, etc.) are now in text_extraction.py
# Note: WhatsAppScraper class is now in whatsapp.py

# API models
class GroupsRequest(BaseModel):
    groups: List[str]


class ChatRequest(BaseModel):
    query: str


class GoogleAuthRequest(BaseModel):
    access_token: Optional[str] = None
    credentials: Optional[dict] = None


# ==================== Google OAuth2 Authentication ====================

@app.get("/auth/google")
async def google_auth():
    """Initiate Google OAuth flow"""
    try:
        if not CLIENT_ID or not CLIENT_SECRET:
            raise HTTPException(
                status_code=500,
                detail="CLIENT_ID and CLIENT_SECRET must be set in .env file"
            )
        
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": CLIENT_ID,
                    "client_secret": CLIENT_SECRET,
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "redirect_uris": [REDIRECT_URI]
                }
            },
            scopes=SCOPES
        )
        flow.redirect_uri = REDIRECT_URI
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent'
        )
        
        # Store state for verification
        credentials_store['state'] = state
        
        return JSONResponse({
            'auth_url': authorization_url,
            'state': state
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'OAuth initialization failed: {str(e)}')

@app.get("/oauth2callback")
async def oauth2callback(request: Request):
    """Handle OAuth2 callback"""
    try:
        authorization_response = str(request.url)
        
        # Log the callback for debugging
        print(f"[INFO] OAuth2 callback received")
        
        # Always use SCOPES constant to ensure both Drive and Gmail are included
        # Don't rely on scopes from URL as they may be incomplete
        print(f"[INFO] Using scopes for token exchange: {SCOPES}")
        
        # Create flow with ALL required scopes
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": CLIENT_ID,
                    "client_secret": CLIENT_SECRET,
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "redirect_uris": [REDIRECT_URI]
                }
            },
            scopes=SCOPES  # Always use SCOPES constant with both Drive and Gmail
        )
        flow.redirect_uri = REDIRECT_URI
        
        # Fetch token with the authorization response
        flow.fetch_token(authorization_response=authorization_response)
        
        credentials = flow.credentials
        
        # Log what scopes were actually returned
        print(f"[INFO] Token received with scopes: {credentials.scopes}")
        
        # Store credentials with ALL required scopes (merge with SCOPES)
        token_key = 'user_token'
        merged_scopes = list(set(credentials.scopes or []) | set(SCOPES))
        credentials_store[token_key] = {
            'token': credentials.token,
            'refresh_token': credentials.refresh_token,
            'token_uri': credentials.token_uri,
            'client_id': credentials.client_id,
            'client_secret': credentials.client_secret,
            'scopes': merged_scopes
        }
        
        # Save token.json for MCP server compatibility
        # This allows the MCP server to use the same authentication
        token_file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'token.json')
        
        # Create token data with merged scopes to ensure Gmail is included
        token_data = {
            'token': credentials.token,
            'refresh_token': credentials.refresh_token,
            'token_uri': credentials.token_uri,
            'client_id': credentials.client_id,
            'client_secret': credentials.client_secret,
            'scopes': merged_scopes  # Use merged scopes
        }
        with open(token_file_path, 'w') as token_file:
            json.dump(token_data, token_file)
        print(f"[INFO] Token saved to {token_file_path} with scopes: {merged_scopes}")
        
        # Encode credentials as JSON for passing to frontend
        creds_data = {
            'token': credentials.token,
            'refresh_token': credentials.refresh_token if credentials.refresh_token else None,
            'token_uri': credentials.token_uri,
            'client_id': credentials.client_id,
            'client_secret': credentials.client_secret,
            'scopes': list(credentials.scopes) if credentials.scopes else []
        }
        creds_json = json.dumps(creds_data)
        
        # Return success page with token
        html_content = f"""
        <html>
            <head><title>Authentication Successful</title></head>
            <body style="font-family: Arial; text-align: center; padding: 50px;">
                <h1>✅ Authentication Successful!</h1>
                <p>You can now close this window and return to the application.</p>
                <p>Access Token: <code>{credentials.token[:20]}...</code></p>
                <script>
                    // Send credentials to parent window if opened in popup
                    if (window.opener) {{
                        const credsData = {creds_json};
                        window.opener.postMessage({{
                            type: 'google_auth_success',
                            access_token: '{credentials.token}',
                            credentials: credsData
                        }}, '*');
                        window.close();
                    }}
                </script>
            </body>
        </html>
        """
        
        return HTMLResponse(content=html_content)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'OAuth callback failed: {str(e)}')

# ==================== Google Drive Functions (Using MCP Client) ====================

@app.get("/api/drive/status")
async def drive_status():
    """
    Backwards-compatible alias for drive auth status.
    Mirrors /api/drive/mcp-status described in drive-mail-connection.md.
    """
    return await get_mcp_status()


@app.get("/api/drive/auth-url")
async def drive_auth_url():
    """
    Backwards-compatible alias that returns the Google OAuth URL.
    Mirrors /auth/google but under /api/drive/auth-url as documented.
    """
    return await google_auth()

@app.post("/api/drive/files")
async def get_drive_files(payload: GoogleAuthRequest):
    """
    Fetch and extract text from Google Drive files using MCP client.
    Files are fed into RAG system and Excel agent.
    """
    global rag_system, excel_agent_system
    
    try:
        # Get credentials from payload to save token.json if needed
        access_token = payload.access_token
        credentials_data = payload.credentials
        
        # If credentials provided, save them to token.json for MCP client
        if credentials_data:
            token_file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'token.json')
            
            # Check if the provided credentials include Gmail scope
            # If not, the user needs to re-authenticate with the new scopes
            provided_scopes = credentials_data.get('scopes', [])
            gmail_scope = 'https://www.googleapis.com/auth/gmail.readonly'
            
            print(f"[INFO] Received credentials with scopes: {provided_scopes}")
            
            # CRITICAL: Check if Gmail scope is present in the ACCESS TOKEN
            # If not, the access token itself doesn't have gmail permission
            # User MUST re-authenticate to get a new token with gmail scope
            if gmail_scope not in provided_scopes:
                print(f"[WARNING] Gmail scope missing from credentials! User must re-authenticate.")
                print(f"[INFO] To fix this: 1) Delete token.json, 2) Clear browser localStorage, 3) Go to /auth/google")
                raise HTTPException(
                    status_code=401,
                    detail="Gmail scope missing. Please re-authenticate with Google to enable email access. "
                           "Clear your browser cache/localStorage and go to /auth/google to get new credentials "
                           "with both Drive and Gmail permissions."
                )
            
            # All required scopes present - save the token
            token_data = {
                'token': credentials_data.get('token') or access_token,
                'refresh_token': credentials_data.get('refresh_token'),
                'token_uri': credentials_data.get('token_uri', 'https://oauth2.googleapis.com/token'),
                'client_id': credentials_data.get('client_id') or CLIENT_ID,
                'client_secret': credentials_data.get('client_secret') or CLIENT_SECRET,
                'scopes': provided_scopes
            }
            
            print(f"[INFO] Saving token with scopes: {provided_scopes}")
            with open(token_file_path, 'w') as token_file:
                json.dump(token_data, token_file)
            print(f"[INFO] Token saved to {token_file_path} for MCP client")
        
        # Get MCP client and reset service cache to use new credentials
        mcp_client = get_mcp_client()
        mcp_client.reset_service()  # Clear cached service to use new credentials
        
        if not mcp_client.is_authenticated():
            raise HTTPException(
                status_code=401, 
                detail="Not authenticated. Please authenticate with Google first via /auth/google"
            )
        
        # Initialize RAG system if not already initialized
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="GROQ_API_KEY not found in environment variables")
        
        if not rag_system:
            rag_system = RAGSystem(api_key)
        
        # Use MCP client to get all files with content
        print("[INFO] Using MCP client to fetch Google Drive files...")
        file_results = mcp_client.get_all_files_with_content(GOOGLE_DRIVE_DOWNLOAD_DIR)
        
        extracted_files = []
        downloaded_file_paths = []
        
        for file_data in file_results:
            if file_data.get('error'):
                extracted_files.append({
                    'id': file_data.get('id'),
                    'name': file_data.get('name'),
                    'type': file_data.get('type'),
                    'size': 'Unknown',
                    'extractedText': file_data.get('extractedText', 'Error processing file'),
                    'error': True
                })
            else:
                extracted_files.append({
                    'id': file_data.get('id'),
                    'name': file_data.get('name'),
                    'type': file_data.get('type'),
                    'size': file_data.get('size'),
                    'modified': file_data.get('modified'),
                    'extractedText': file_data.get('extractedText'),
                })
                if file_data.get('path'):
                    downloaded_file_paths.append(file_data['path'])
        
        # Separate Excel files from other files
        excel_file_paths = [p for p in downloaded_file_paths if p.lower().endswith(('.xlsx', '.xls'))]
        
        # Load Excel files into Excel Agent (not RAG)
        excel_agents_created = 0
        if excel_file_paths and LANGCHAIN_AGENT_AVAILABLE:
            if not excel_agent_system:
                excel_agent_system = ExcelAgentSystem(api_key)
            
            for excel_path in excel_file_paths:
                if excel_agent_system.add_excel_file(excel_path):
                    excel_agents_created += 1
                    print(f"[INFO] Excel agent created for Google Drive file: {os.path.basename(excel_path)}")
        
        # Load non-Excel files into RAG system only
        num_chunks = 0
        if downloaded_file_paths:
            try:
                documents = rag_system.load_documents(downloaded_file_paths)
                
                if documents:
                    num_chunks = rag_system.create_vector_store(documents)
                    print(f"[DEBUG] Loaded {len(documents)} non-Excel documents from Google Drive into RAG system, created {num_chunks} chunks")
                else:
                    print("[INFO] No non-Excel documents could be loaded from Google Drive files")
            except Exception as e:
                print(f"[ERROR] Error loading Google Drive files into RAG system: {str(e)}")
                traceback.print_exc()
        
        # ==================== Gmail Email Fetching ====================
        # Fetch latest 5 emails from Gmail and load into RAG system
        email_count = 0
        extracted_emails = []
        try:
            print("[INFO] Fetching latest 5 emails from Gmail...")
            # Use the same mcp_client for Gmail (it now handles both Drive and Gmail)
            mcp_client.reset_service()  # Clear cached service to use new credentials
            
            # Fetch emails formatted for RAG
            emails = mcp_client.get_emails_for_rag(max_results=5)
            
            if emails:
                # Store email data for response
                for email in emails:
                    extracted_emails.append({
                        'id': email.get('id'),
                        'subject': email.get('subject'),
                        'from': email.get('from'),
                        'date': email.get('date'),
                        'extractedText': email.get('extractedText'),
                        'type': 'email'
                    })
                
                # Load emails into RAG system
                email_docs = rag_system.load_gmail_emails(emails)
                if email_docs:
                    email_chunks = rag_system.create_vector_store(email_docs)
                    email_count = len(emails)
                    print(f"[INFO] Successfully loaded {email_count} emails into RAG system with {email_chunks} chunks")
                else:
                    print("[INFO] No email documents could be created")
            else:
                print("[INFO] No emails found in Gmail inbox")
                
        except Exception as e:
            print(f"[WARNING] Error fetching Gmail emails: {str(e)}")
            print("[INFO] Gmail integration failed - continuing with Drive files only")
            traceback.print_exc()
        
        return JSONResponse({
            'files': extracted_files,
            'total': len(extracted_files),
            'loaded_to_rag': num_chunks,
            'excel_agents_created': excel_agents_created,
            'excel_agent_available': excel_agent_system is not None and len(excel_agent_system.agents) > 0 if excel_agent_system else False,
            'emails': extracted_emails,
            'emails_loaded': email_count,
            'message': f'Successfully processed {len(extracted_files)} files from Google Drive and {email_count} emails from Gmail',
            'mcp_enabled': True,
            'gmail_enabled': True
        })
        
    except HTTPException:
        raise
    except Exception as e:
        error_details = traceback.format_exc()
        print(f"Error in get_drive_files: {error_details}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/drive/files")
async def get_drive_files_using_saved_token():
    """
    GET variant of /api/drive/files for compatibility with the documented API.
    Uses the saved token.json (no request body needed).
    """
    empty_payload = GoogleAuthRequest(access_token=None, credentials=None)
    return await get_drive_files(empty_payload)


@app.get("/api/drive/mcp-status")
async def get_mcp_status():
    """Check MCP client authentication status for Google Drive"""
    try:
        mcp_client = get_mcp_client()
        is_authenticated = mcp_client.is_authenticated()
        
        return JSONResponse({
            'mcp_enabled': True,
            'authenticated': is_authenticated,
            'token_file_exists': is_authenticated,
            'message': 'Authenticated with Google Drive' if is_authenticated else 'Not authenticated. Use /auth/google to authenticate.'
        })
    except Exception as e:
        return JSONResponse({
            'mcp_enabled': True,
            'authenticated': False,
            'error': str(e)
        })


@app.get("/api/emails")
async def get_gmail_emails(max_results: int = 5):
    """
    Fetch Gmail emails with AI-generated summaries.
    Returns emails formatted for the EmailManager component.
    """
    try:
        mcp_client = get_mcp_client()
        
        if not mcp_client.is_authenticated():
            # Return empty list instead of 401 to handle UI gracefully
            return JSONResponse({
                'success': False,
                'emails': [],
                'message': 'Not authenticated'
            })
        
        # Fetch emails from Gmail
        print(f"[API] Fetching {max_results} emails from Gmail...")
        raw_emails = mcp_client.fetch_latest_emails(max_results)
        if not raw_emails:
            print("[API] No emails returned from Gmail; using fallback sample emails for UI demo.")
            raw_emails = [
                {
                    "id": "demo-1",
                    "subject": "Welcome to TOAI Email Manager",
                    "from": "TOAI Demo <demo@toai.ai>",
                    "to": "user@example.com",
                    "date": "Tue, 17 Dec 2024 10:00:00 +0000",
                    "body": "This is a demo email used when no Gmail messages are available.",
                    "snippet": "This is a demo email used when no Gmail messages are available.",
                    "threadId": "demo-thread-1",
                    "messageId": "<demo-1@toai.ai>",
                },
                {
                    "id": "demo-2",
                    "subject": "Connect Gmail to see your real emails",
                    "from": "Support <support@toai.ai>",
                    "to": "user@example.com",
                    "date": "Tue, 17 Dec 2024 11:00:00 +0000",
                    "body": "TOAI could not find any recent messages in your Gmail inbox. "
                    "Once you send or receive emails, they will start appearing here.",
                    "snippet": "TOAI could not find any recent messages in your Gmail inbox…",
                    "threadId": "demo-thread-2",
                    "messageId": "<demo-2@toai.ai>",
                },
            ]
        
        # Generative AI for summaries
        api_key = os.getenv("GROQ_API_KEY")
        formatted_emails = []
        
        for email in raw_emails:
            try:
                # Generate AI summary using Groq
                summary = ""
                if api_key and email.get('body'):
                    try:
                        from groq import Groq
                        client = Groq(api_key=api_key)
                        
                        # Create summary prompt
                        body_text = email['body'][:3000]  # Limit context window
                        summary_prompt = f"""Summarize this email in 2 sentences. Focus on purpose and action items:

Subject: {email['subject']}
From: {email['from']}
Content: {body_text}

Summary:"""
                        
                        completion = client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=[{"role": "user", "content": summary_prompt}],
                            max_tokens=100,
                            temperature=0.3
                        )
                        summary = completion.choices[0].message.content.strip()
                    except Exception as e:
                        print(f"[API] Summary generation failed: {e}")
                        summary = email.get('snippet', '')
                else:
                    summary = email.get('snippet', '')
                
                # Parse sender
                from_str = email.get('from', 'Unknown')
                sender_name = from_str
                sender_email = from_str
                
                # Simple parsing for "Name <email>"
                if '<' in from_str and '>' in from_str:
                    parts = from_str.split('<')
                    sender_name = parts[0].strip().strip('"')
                    sender_email = parts[1].strip('>')
                
                # Parse date to ISO string
                date_str = email.get('date', '')
                timestamp = date_str # Default to raw string if parsing fails
                try:
                    from email.utils import parsedate_to_datetime
                    dt = parsedate_to_datetime(date_str)
                    timestamp = dt.isoformat()
                except:
                    pass
                
                formatted_emails.append({
                    'id': email['id'],
                    'from': sender_name,
                    'fromEmail': sender_email,
                    'to': email.get('to', ''),
                    'subject': email['subject'],
                    'preview': email.get('snippet', '')[:100] + '...',
                    'body': email['body'],
                    'summary': summary,
                    'timestamp': timestamp,
                    'read': False, # New emails are unread
                    'category': 'Internal' if 'toai' in sender_email.lower() else 'Client', # Simple heuristic
                    'threadId': email.get('threadId'),  # Gmail thread ID for replying in thread
                    'messageId': email.get('messageId')  # Email Message-ID for In-Reply-To header
                })
                
            except Exception as e:
                print(f"[API] Error formatting email {email.get('id')}: {e}")
                continue
        
        return JSONResponse({'success': True, 'emails': formatted_emails})
        
    except Exception as e:
        print(f"[API] Error in get_gmail_emails: {e}")
        return JSONResponse({
            'success': False,
            'error': str(e)
        }, status_code=500)


class EmailReplyRequest(BaseModel):
    original_body: str
    original_sender: str
    original_subject: str
    user_prompt: Optional[str] = None
    file_context: Optional[str] = None

class EmailSendRequest(BaseModel):
    to: str
    subject: str
    body: str
    cc: Optional[str] = None
    thread_id: Optional[str] = None  # Gmail thread ID for replying in thread
    message_id: Optional[str] = None  # Original message ID for In-Reply-To header

@app.post("/api/email/generate")
async def generate_email_reply(request: EmailReplyRequest):
    """Generate an AI-powered email reply based on the original email and user prompt."""
    try:
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            return JSONResponse(
                status_code=500,
                content={"success": False, "error": "GROQ_API_KEY not configured"}
            )
        
        from groq import Groq
        client = Groq(api_key=api_key)
        
        # Build the context for the AI
        context_parts = [
            f"Original Email from {request.original_sender}:",
            f"Subject: {request.original_subject}",
            f"Content: {request.original_body[:2000]}"
        ]
        
        # Add file context if provided
        if request.file_context:
            context_parts.append(f"\n--- Attached Documents for Reference ---\n{request.file_context[:3000]}")
        
        context = "\n".join(context_parts)
        
        # Build the prompt
        if request.user_prompt:
            user_instruction = f"User's instructions: {request.user_prompt}"
        else:
            user_instruction = "Generate a polite and professional reply to this email."
        
        system_prompt = """You are a professional email assistant. Generate ONLY the email reply text - nothing else.

IMPORTANT RULES:
- Do NOT include any introductory text like "Here's a draft reply" or "Here's my response"
- Do NOT include explanations of what you're doing
- Do NOT include meta-commentary about the email
- Start DIRECTLY with the greeting (e.g., "Dear...", "Hi...", "Hello...")
- End with a professional closing and signature placeholder like "[Your Name]"
- Keep the reply professional, clear, and address the key points from the original email
- If the user provides attached documents, incorporate relevant information from them
- Do not include email headers like "From:", "To:", "Subject:"

Output only the email body text, nothing else."""
        
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"{context}\n\n{user_instruction}"}
            ],
            max_tokens=800,
            temperature=0.7
        )
        
        reply = response.choices[0].message.content.strip()
        return {"success": True, "reply": reply}
        
    except Exception as e:
        print(f"[ERROR] Email generation failed: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )

@app.post("/api/email/send")
async def send_email(request: EmailSendRequest):
    """Send an email using Gmail API."""
    try:
        mcp_client = get_mcp_client()
        
        if not mcp_client.is_authenticated():
            return JSONResponse(
                status_code=401,
                content={"success": False, "error": "Not authenticated. Please sign in with Google first."}
            )
        
        # Debug logging
        print(f"[API] Sending email reply:")
        print(f"  To: {request.to}")
        print(f"  Subject: {request.subject}")
        print(f"  Thread ID: {request.thread_id}")
        print(f"  Message ID: {request.message_id}")
        
        result = mcp_client.send_email(
            to=request.to,
            subject=request.subject,
            body=request.body,
            cc=request.cc,
            thread_id=request.thread_id,
            message_id=request.message_id
        )
        
        if result.get("success"):
            return {"success": True, "message": "Email sent successfully", "message_id": result.get("message_id")}
        else:
            return JSONResponse(
                status_code=500,
                content={"success": False, "error": result.get("error", "Failed to send email")}
            )
            
    except Exception as e:
        print(f"[ERROR] Email sending failed: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


# ==================== WhatsApp Integration ====================

@app.post("/api/whatsapp/connect")
def connect_whatsapp():
    """Start WhatsApp Web connection (opens browser with QR code)."""
    global whatsapp_driver, whatsapp_connected
    
    try:
        print("[WhatsApp] Starting WhatsApp Web connection...")
        scraper = WhatsAppScraper(download_dir=WHATSAPP_DOWNLOAD_DIR, images_dir=WHATSAPP_IMAGES_DIR)
        success = scraper.start_whatsapp_async()  # Non-blocking version
        
        if success:
            whatsapp_driver = scraper
            whatsapp_connected = False  # Not logged in yet, waiting for QR scan
            print("[WhatsApp] Browser opened, waiting for QR code scan...")
            return {"message": "WhatsApp Web opened, please scan QR code", "status": "waiting_for_scan"}
        else:
            raise HTTPException(status_code=500, detail="Failed to start WhatsApp Web")
    except Exception as e:
        print(f"[WhatsApp] Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/whatsapp/qr")
def get_whatsapp_qr():
    """Get the WhatsApp QR code as base64 image."""
    global whatsapp_driver
    
    if not whatsapp_driver:
        raise HTTPException(status_code=400, detail="WhatsApp not initialized. Call /api/whatsapp/connect first.")
    
    try:
        # Check if already logged in
        if whatsapp_driver.is_logged_in():
            return {"logged_in": True, "qr_code": None}
        
        # Get QR code
        qr_base64 = whatsapp_driver.get_qr_code_base64()
        
        if qr_base64:
            return {"logged_in": False, "qr_code": qr_base64}
        else:
            return {"logged_in": False, "qr_code": None, "message": "QR code not found, may need to refresh"}
    except Exception as e:
        print(f"[WhatsApp] Error getting QR: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/whatsapp/login-status")
def get_whatsapp_login_status():
    """Check if WhatsApp is logged in after QR scan."""
    global whatsapp_driver, whatsapp_connected
    
    if not whatsapp_driver:
        return {"logged_in": False, "message": "WhatsApp not initialized"}
    
    try:
        is_logged_in = whatsapp_driver.is_logged_in()
        if is_logged_in:
            whatsapp_connected = True
        return {"logged_in": is_logged_in}
    except Exception as e:
        print(f"[WhatsApp] Error checking login status: {str(e)}")
        return {"logged_in": False, "error": str(e)}


@app.get("/api/whatsapp/groups")
def get_whatsapp_groups():
    """Return a list of visible WhatsApp chats/groups from the left pane."""
    global whatsapp_driver, whatsapp_connected

    if not whatsapp_driver or not whatsapp_connected:
        raise HTTPException(status_code=400, detail="WhatsApp not connected. Please scan QR code first.")

    try:
        groups = whatsapp_driver.list_groups(limit=50)
        return {"groups": groups}
    except Exception as e:
        print(f"[WhatsApp] Error listing groups: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/whatsapp/scrape")
def scrape_whatsapp(payload: GroupsRequest):
    """Scrape messages from WhatsApp groups."""
    global whatsapp_driver, rag_system
    
    if not whatsapp_driver or not whatsapp_connected:
        raise HTTPException(status_code=400, detail="WhatsApp not connected. Please scan QR code first.")
    
    groups = payload.groups or []
    
    if not groups:
        raise HTTPException(status_code=400, detail="No groups selected")
    
    try:
        api_key = os.getenv("GROQ_API_KEY")
        if not rag_system:
            rag_system = RAGSystem(api_key)
        
        all_messages = []
        all_pdfs = []
        all_ocr_texts = []
        
        for group_name in groups:
            print(f"[WhatsApp] Opening group: {group_name}")
            success = whatsapp_driver.open_group(group_name)
            if success:
                # Extract messages
                messages = whatsapp_driver.extract_messages()
                all_messages.extend(messages)
                print(f"[WhatsApp] Extracted {len(messages)} messages from {group_name}")
                
                # Download PDFs
                pdfs = whatsapp_driver.download_pdfs()
                all_pdfs.extend(pdfs)
                print(f"[WhatsApp] Downloaded {len(pdfs)} PDFs from {group_name}")
                
                # Download images and perform OCR
                ocr_results = whatsapp_driver.download_images_and_ocr()
                all_ocr_texts.extend(ocr_results)
                print(f"[WhatsApp] Processed {len(ocr_results)} images from {group_name}")
            else:
                print(f"[WhatsApp] Could not open group: {group_name}")
        
        # Load messages into RAG
        if all_messages:
            msg_docs = rag_system.load_whatsapp_messages(all_messages)
            rag_system.create_vector_store(msg_docs)
            print(f"[WhatsApp] Loaded {len(all_messages)} messages into RAG system")
        
        # Load PDFs into RAG
        if all_pdfs:
            pdf_docs = rag_system.load_documents(all_pdfs)
            if pdf_docs:
                rag_system.create_vector_store(pdf_docs)
                print(f"[WhatsApp] Loaded {len(all_pdfs)} PDFs into RAG system")
        
        # Load OCR texts from images into RAG
        if all_ocr_texts:
            ocr_docs = rag_system.load_ocr_texts(all_ocr_texts)
            if ocr_docs:
                rag_system.create_vector_store(ocr_docs)
                print(f"[WhatsApp] Loaded {len(all_ocr_texts)} OCR results into RAG system")
        
        return {
            "message": "WhatsApp data scraped successfully",
            "messages": len(all_messages),
            "pdfs": len(all_pdfs),
            "images_processed": len(all_ocr_texts),
            "images_with_text": sum(1 for x in all_ocr_texts if x.get('text', '').strip()),
        }
    except Exception as e:
        print(f"[WhatsApp] Scraping error: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/whatsapp/disconnect")
def disconnect_whatsapp():
    """Disconnect WhatsApp and close browser."""
    global whatsapp_driver, whatsapp_connected
    
    try:
        if whatsapp_driver:
            whatsapp_driver.close()
            whatsapp_driver = None
        whatsapp_connected = False
        print("[WhatsApp] Disconnected successfully")
        return {"message": "WhatsApp disconnected"}
    except Exception as e:
        print(f"[WhatsApp] Error disconnecting: {str(e)}")
        # Still mark as disconnected even if close fails
        whatsapp_driver = None
        whatsapp_connected = False
        return {"message": "WhatsApp disconnected (with errors)", "error": str(e)}


# API Routes - File Upload
@app.post("/api/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    """Upload files to the RAG system."""
    global rag_system, excel_agent_system
    
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    saved_paths = []
    excel_paths = []
    
    for file in files:
        if file.filename:
            filepath = os.path.join(UPLOAD_FOLDER, file.filename)
            try:
                content = await file.read()
                with open(filepath, "wb") as f:
                    f.write(content)
                saved_paths.append(filepath)
                
                # Track Excel files separately
                if file.filename.lower().endswith(('.xlsx', '.xls')):
                    excel_paths.append(filepath)
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error saving file {file.filename}: {str(e)}")
    
    try:
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="GROQ_API_KEY not found in environment variables")
        
        if not rag_system:
            rag_system = RAGSystem(api_key)
        
        # Initialize Excel agent system and load Excel files
        excel_agents_created = 0
        if excel_paths and LANGCHAIN_AGENT_AVAILABLE:
            if not excel_agent_system:
                excel_agent_system = ExcelAgentSystem(api_key)
            
            for excel_path in excel_paths:
                if excel_agent_system.add_excel_file(excel_path):
                    excel_agents_created += 1
                    print(f"[INFO] Excel agent created for: {os.path.basename(excel_path)}")
        
        # Load non-Excel documents into RAG system (Excel files go to Excel Agent only)
        non_excel_paths = [p for p in saved_paths if not p.lower().endswith(('.xlsx', '.xls'))]
        documents = rag_system.load_documents(non_excel_paths)
        
        num_chunks = 0
        if documents:
            num_chunks = rag_system.create_vector_store(documents)
            print(f"[INFO] Loaded {len(documents)} documents into RAG system")
        
        return {
            "message": f"Successfully processed {len(files)} files",
            "chunks": num_chunks,
            "excel_agents_created": excel_agents_created,
            "excel_agent_available": excel_agent_system is not None and len(excel_agent_system.agents) > 0 if excel_agent_system else False,
        }
    except HTTPException:
        raise
    except Exception as e:
        error_detail = f"Error processing files: {str(e)}\n{traceback.format_exc()}"
        raise HTTPException(status_code=500, detail=error_detail)


# API Routes - Chat
@app.post("/api/chat")
def chat(payload: ChatRequest):
    """Process chat query using RAG system and/or Excel agent."""
    global rag_system, excel_agent_system, agentic_router
    
    query = payload.query or ""
    
    if not query:
        raise HTTPException(status_code=400, detail="No query provided")
    
    try:
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="GROQ_API_KEY not found in environment variables")
        
        if not rag_system:
            rag_system = RAGSystem(api_key)
        
        response = ""
        query_type = "general"
        tools_used = []
        
        # Check if we have data sources
        has_rag = rag_system and rag_system.vector_store is not None
        has_excel = excel_agent_system is not None and bool(excel_agent_system.agents)
        
        # Route query based on available tools
        if has_excel and excel_agent_system.is_excel_query(query):
            # Use Excel agent for quantitative queries - DO NOT mix with RAG
            query_type = "excel"
            tools_used = ["Excel_Data_Analyst"]
            excel_analysis = excel_agent_system.run_query(query)
            
            # Get the Excel file names that were used
            excel_files = list(excel_agent_system.dataframes.keys())
            excel_sources = ", ".join(excel_files) if excel_files else "Excel data"
            
            # Generate response with ONLY Excel sources - no RAG context
            combined_prompt = f"""Answer the user's question using ONLY the Excel data analysis result below.

User Question: {query}

Excel Data Analysis Result:
{excel_analysis}

IMPORTANT: 
- Base your answer ONLY on the Excel data analysis above
- Do NOT reference any other documents
- At the end, include a **Sources:** section listing ONLY the Excel file(s): {excel_sources}

Provide a clear, concise answer based on the Excel data."""

            messages = [{
                "role": "system", 
                "content": f"""You are a helpful data analyst. You are answering a question using Excel data.
                
CRITICAL RULES:
1. ONLY cite Excel file sources: {excel_sources}
2. Do NOT mention or cite any PDF, DOCX, or other document files
3. Your answer is based purely on Excel data analysis
4. End with **Sources:** section listing only the Excel file(s) used"""
            }]
            messages.append({"role": "user", "content": combined_prompt})
            
            chat_completion = rag_system.groq_client.chat.completions.create(
                messages=messages,
                model="llama-3.3-70b-versatile",
                temperature=0.3,
                max_tokens=2048,
            )
            response = chat_completion.choices[0].message.content
            
        elif has_rag:
            # Use RAG for document-based queries
            query_type = "rag"
            tools_used = ["PDF_Document_Knowledge_Base"]
            context, sources = rag_system.retrieve_context(query, k=10)
            response = rag_system.generate_response(query, context, sources)
            
        else:
            # No data sources - use general knowledge
            query_type = "general"
            messages = [
                {"role": "system", "content": "You are a helpful assistant. Be informative and helpful."}
            ]
            messages.append({"role": "user", "content": query})
            
            chat_completion = rag_system.groq_client.chat.completions.create(
                messages=messages,
                model="llama-3.3-70b-versatile",
                temperature=0.3,
                max_tokens=2048,
            )
            response = chat_completion.choices[0].message.content
            rag_system.chat_history.append({"role": "user", "content": query})
            rag_system.chat_history.append({"role": "assistant", "content": response})
        
        # Format tables in response
        formatted_response = format_tables_in_response(response)
        
        return {
            "response": formatted_response,
            "query_type": query_type,
            "tools_used": tools_used,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/status")
def status():
    """Get the current status of all systems."""
    return {
        "whatsapp_connected": whatsapp_connected,
        "rag_initialized": rag_system is not None and rag_system.vector_store is not None,
        "excel_agent_initialized": excel_agent_system is not None and len(excel_agent_system.agents) > 0 if excel_agent_system else False,
        "excel_files_loaded": list(excel_agent_system.dataframes.keys()) if excel_agent_system and excel_agent_system.dataframes else [],
    }


@app.post("/api/reset")
def reset_rag():
    """Reset the in-memory RAG system, Excel agent, agentic router, and clear all uploaded files."""
    global rag_system, excel_agent_system, agentic_router
    
    # Reset in-memory systems
    rag_system = None
    excel_agent_system = None
    agentic_router = None
    
    # Clear uploaded files
    cleared_counts = {
        "uploads": 0,
        "whatsapp_downloads": 0,
        "whatsapp_images": 0,
        "google_drive_downloads": 0
    }
    
    directories_to_clear = [
        (UPLOAD_FOLDER, "uploads"),
        (WHATSAPP_DOWNLOAD_DIR, "whatsapp_downloads"),
        (WHATSAPP_IMAGES_DIR, "whatsapp_images"),
        (GOOGLE_DRIVE_DOWNLOAD_DIR, "google_drive_downloads")
    ]
    
    for dir_path, name in directories_to_clear:
        if os.path.exists(dir_path):
            try:
                for item in os.listdir(dir_path):
                    item_path = os.path.join(dir_path, item)
                    try:
                        if os.path.isfile(item_path):
                            os.remove(item_path)
                            cleared_counts[name] += 1
                        elif os.path.isdir(item_path):
                            shutil.rmtree(item_path)
                            cleared_counts[name] += 1
                    except Exception as e:
                        print(f"[RESET] Could not remove {item_path}: {e}")
            except Exception as e:
                print(f"[RESET] Error clearing {dir_path}: {e}")
    
    print(f"[RESET] Cleared: {cleared_counts}")
    return {
        "message": "RAG system, Excel agent, agentic router, and all uploaded files reset successfully",
        "files_cleared": cleared_counts
    }


@app.exception_handler(Exception)
async def global_exception_handler(request, exc: Exception):
    if isinstance(exc, HTTPException):
        return JSONResponse(status_code=exc.status_code, content={"error": exc.detail})
    return JSONResponse(status_code=500, content={"error": str(exc)})


if __name__ == "__main__":
    import uvicorn

    # Using port 5001 to avoid conflict with Express server on 5000
    uvicorn.run("app:app", host="0.0.0.0", port=5001, reload=True)


